"""Smart Spec Generator — Sonnet reads .sav + questionnaire + ticket → TabulateSpec.

This is the "AI data processor" that replaces manual configuration.
Sonnet receives all available context and generates a complete tabulation spec
with confidence levels and questions for the user when uncertain.

Supports:
- Questionnaire: .docx or .pdf (extracts text, detects skip logic, routing)
- Reporting Ticket: .docx (extracts banners, stubs, nets, sig level)
- SPSS metadata: variables, labels, value_labels, n_valid, detected_groups
- Inferred logic: if Q12 has 30 valid out of 85, and label says "PsO", Sonnet infers filter

Cost: ~$0.03-0.05 per call (Sonnet with ~8K input tokens)
"""

import io
import json
import logging
from typing import Any

from config import get_settings

logger = logging.getLogger(__name__)

SMART_SPEC_PROMPT = """You are an expert market research data processor. Your job is to generate a complete tabulation specification from three inputs:

1. **SPSS metadata** — variable names, labels, value labels, response counts, detected groups
2. **Questionnaire** — the survey instrument with question text, response options, skip logic, routing
3. **Reporting Ticket** — client instructions specifying what tables to produce

## Your output

Generate a JSON spec that our tabulation engine can execute directly. Be precise about variable matching.

## Rules

### Variable matching
- Match questionnaire Q-numbers (Q1, Q2, S2, S4A) to actual SPSS variable names using labels
- The questionnaire may number questions differently than the SPSS file
- Use the variable LABEL to match, not just the name

### Skip logic / conditional questions
- If a question was only asked to a subgroup (e.g., "PsO patients only"), detect this from:
  a) The questionnaire text ("treatments for your Plaque Psoriasis")
  b) The n_valid count (if 30 out of 85 responded, it's likely conditional)
  c) The variable label containing a condition name
- For conditional questions, add a `filter` field with the condition

### Banners
- Extract from the reporting ticket ("Generally, show data by: PsO/PsA/UC-CD")
- Map to actual SPSS variables
- If ticket says to combine categories (e.g., "UC/CD Patients"), create a custom_group

### Nested banners
- Some questions show results by TWO dimensions (e.g., "by treatment concept AND by condition")
- Detect from ticket instructions like "show by Long-Acting Injectable / Once Weekly Oral"
- Create nested_banner entries: {parent_var, child_var}

### Nets
- T2B (Top 2 Box): top 2 values on the scale
- B2B (Bottom 2 Box): bottom 2 values
- M3B (Middle 3 Box): middle values
- For 7-point scales: T2B=[6,7], M3B=[3,4,5], B2B=[1,2]
- For 5-point scales: T2B=[4,5], B2B=[1,2]
- Include the scale endpoint labels in net names when available

### MRS (Multiple Response Sets)
- Binary yes/no variables with shared question prefix → group as MRS
- The label suffix after "-" or ":" is the option name
- Base = respondents who answered at least one option (NOT total sample)

### Label Enrichment (when questionnaire is provided)
- Compare SPSS variable labels with questionnaire text
- If an SPSS label is truncated (e.g., "Conc" instead of "Concerns with how well the product works"),
  provide the full text from the questionnaire in label_overrides
- Only include overrides where the SPSS label is clearly truncated or incomplete
- The user can choose whether to apply these overrides or keep original labels
- Format: {"variable_name": "Full text from questionnaire"}

### Confidence
- Rate each decision: "high" (clear from text), "medium" (inferred), "low" (guessing)
- If confidence is low on a critical decision, add it to questions_for_user

## Output format

```json
{
  "title": "Study title",
  "banners": ["variable_name"],
  "custom_groups": [
    {"name": "UC/CD Patients", "conditions": [{"variable": "Q_10", "operator": "in", "values": [3, 4]}]}
  ],
  "stubs": ["var1", "var2"],
  "stub_filters": {
    "Q_12_1": {"variable": "Q_10", "operator": "eq", "value": 1},
    "Q_13_1": {"variable": "Q_10", "operator": "eq", "value": 2}
  },
  "mrs_groups": {
    "Treatments PsO": ["Q_12_1", "Q_12_2", "Q_12_3"]
  },
  "grid_groups": {
    "Oral Treatment Satisfaction": {"variables": ["Q_20_1", "Q_20_2", "Q_20_3"], "show": ["t2b", "b2b", "mean"]}
  },
  "nested_banners": [
    {"parent_var": "concept_type", "child_var": "Q_10", "parent_labels": {"injectable": "Long-Acting Injectable", "oral": "Once Weekly Oral"}}
  ],
  "nets": {
    "Q_27_1": {"Top 2 Box (6-7 Extremely appealing)": [6, 7], "Middle 3 Box (3-5)": [3, 4, 5], "Bottom 2 Box (1-2 Not appealing)": [1, 2]}
  },
  "significance_level": 0.95,
  "weight": null,
  "include_means": true,
  "include_total_column": true,
  "confidence": 0.85,
  "decisions": [
    {"decision": "Q12 filtered to PsO patients only", "confidence": "high", "reason": "Label says 'PsO', 30/85 answered"},
    {"decision": "UC and CD combined as one banner group", "confidence": "medium", "reason": "Ticket says 'UC/CD Patients'"}
  ],
  "questions_for_user": [
    "Q27 mentions two treatment concepts. Should I create a nested banner (Injectable x Condition, Oral x Condition)?",
    "Q14 treatments: should UC and CD patients be shown separately or combined?"
  ],
  "matched_variables": {
    "S2": {"matched_to": "Q_10", "label": "Which of the following...", "confidence": "high"},
    "S4A": {"matched_to": "Q_12_1..Q_12_20", "label": "Treatments PsO", "confidence": "high"}
  },
  "label_overrides": {
    "Q_37_1": "Concerns with how well the product works",
    "Q_37_2": "Concerns with safety",
    "Q_37_3": "Personal discomfort with dosing frequency"
  }
}
```

Return ONLY valid JSON. No markdown, no explanation outside the JSON."""


class SmartSpecGenerator:
    """Generate complete TabulateSpec from .sav + questionnaire + ticket using Sonnet."""

    def __init__(self):
        settings = get_settings()
        if not settings.anthropic_api_key:
            raise ValueError("ANTHROPIC_API_KEY required for smart spec generation")
        from anthropic import AsyncAnthropic
        self.client = AsyncAnthropic(api_key=settings.anthropic_api_key)

    async def generate(
        self,
        metadata: dict[str, Any],
        questionnaire_text: str | None = None,
        ticket_text: str | None = None,
    ) -> dict[str, Any]:
        """Generate a complete TabulateSpec from available context.

        Args:
            metadata: Output of QuantiProEngine.extract_metadata()
            questionnaire_text: Extracted text from questionnaire .docx/.pdf
            ticket_text: Extracted text from reporting ticket .docx

        Returns:
            Dict with spec + confidence + decisions + questions_for_user
        """
        # Build compact metadata summary for Sonnet
        meta_summary = self._build_metadata_summary(metadata)

        user_content = f"<spss_metadata>\n{meta_summary}\n</spss_metadata>"

        if questionnaire_text:
            user_content += f"\n\n<questionnaire>\n{questionnaire_text[:12000]}\n</questionnaire>"

        if ticket_text:
            user_content += f"\n\n<reporting_ticket>\n{ticket_text[:8000]}\n</reporting_ticket>"

        user_content += "\n\nGenerate the complete tabulation spec. Match every question in the ticket to actual SPSS variables. Detect skip logic, conditional bases, MRS groups, and nested banners. Output ONLY the JSON spec."

        response = await self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8192,
            system=SMART_SPEC_PROMPT,
            messages=[{"role": "user", "content": user_content}],
        )

        response_text = response.content[0].text if response.content else ""
        spec = self._parse_json(response_text)

        # Ensure required fields
        spec.setdefault("banners", [])
        spec.setdefault("stubs", [])
        spec.setdefault("stub_filters", {})
        spec.setdefault("mrs_groups", {})
        spec.setdefault("grid_groups", {})
        spec.setdefault("nested_banners", [])
        spec.setdefault("nets", {})
        spec.setdefault("custom_groups", [])
        spec.setdefault("significance_level", 0.95)
        spec.setdefault("weight", None)
        spec.setdefault("include_means", True)
        spec.setdefault("include_total_column", True)
        spec.setdefault("confidence", 0.5)
        spec.setdefault("decisions", [])
        spec.setdefault("questions_for_user", [])
        spec.setdefault("matched_variables", {})
        spec.setdefault("label_overrides", {})
        spec.setdefault("title", "")

        n_stubs = len(spec["stubs"])
        n_mrs = len(spec["mrs_groups"])
        n_grids = len(spec["grid_groups"])
        n_nested = len(spec["nested_banners"])
        n_questions = len(spec["questions_for_user"])
        logger.info(
            "[SMART_SPEC] Generated: %d stubs, %d MRS, %d grids, %d nested banners, "
            "confidence=%.2f, %d questions for user",
            n_stubs, n_mrs, n_grids, n_nested, spec["confidence"], n_questions,
        )

        return spec

    @staticmethod
    def _build_metadata_summary(metadata: dict) -> str:
        """Build a compact metadata summary for Sonnet's context window."""
        lines = [
            f"File: {metadata.get('file_name', '?')}",
            f"Cases: {metadata.get('n_cases', '?')}",
            f"Variables: {metadata.get('n_variables', '?')}",
            "",
            "VARIABLES (name | label | type | n_valid | n_labels | value_labels_preview):",
        ]

        for v in metadata.get("variables", [])[:200]:
            vl = v.get("value_labels") or {}
            vl_preview = list(vl.values())[:4] if vl else []
            n_valid = v.get("n_valid", "?")
            n_missing = v.get("n_missing", 0)
            lines.append(
                f"  {v['name']:15s} | {(v.get('label') or '')[:80]:80s} | {v.get('type','?'):8s} | "
                f"valid={n_valid} miss={n_missing} | labels={len(vl)} {vl_preview}"
            )

        # Detected groups
        groups = metadata.get("detected_groups") or []
        if groups:
            lines.append("")
            lines.append("DETECTED GROUPS:")
            for g in groups:
                lines.append(
                    f"  [{g.get('question_type','?'):12s}] {g.get('display_name','')[:50]} "
                    f"({len(g.get('variables',[]))} vars: {g.get('variables',[])})"
                )

        # Suggested banners
        banners = metadata.get("suggested_banners") or []
        if banners:
            lines.append("")
            lines.append("SUGGESTED BANNERS:")
            for b in banners:
                lines.append(f"  {b.get('variable','?')} — {b.get('label','')}")

        # Weights
        weights = metadata.get("detected_weights") or []
        if weights:
            lines.append(f"\nDETECTED WEIGHTS: {weights}")

        return "\n".join(lines)

    @staticmethod
    def _parse_json(text: str) -> dict:
        """Extract JSON from Sonnet's response."""
        text = text.strip()
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            start = text.find("{")
            end = text.rfind("}") + 1
            if start >= 0 and end > start:
                try:
                    return json.loads(text[start:end])
                except json.JSONDecodeError:
                    pass
            return {"decisions": [{"decision": "Failed to parse Sonnet response", "confidence": "low"}]}

    @staticmethod
    def extract_document_text(file_bytes: bytes, filename: str) -> str:
        """Extract text from .docx or .pdf file."""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext in ("docx", "doc"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)

        elif ext == "pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
            except ImportError:
                # Fallback: try pdfplumber
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        return "\n".join(page.extract_text() or "" for page in pdf.pages)
                except ImportError:
                    return "[PDF parsing not available — install PyMuPDF or pdfplumber]"

        return f"[Unsupported format: .{ext}]"
