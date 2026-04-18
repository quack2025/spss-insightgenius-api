"""Parse Reporting Ticket .docx files into TabulateSpec via Claude Sonnet.

Sonnet is used instead of Haiku because:
1. Variable fuzzy matching: ticket says "S2" but .sav has "Q_10" — needs reasoning
2. Custom group logic: "UC/CD Patients" = combine two conditions
3. Net interpretation: "B2B, M3B, T2B" on different scales
Cost: ~$0.01 per ticket (vs $0.002 for Haiku). Worth it — a bad parse wastes the user's time.
"""

import io
import json
import logging
from typing import Any

from config import get_settings

logger = logging.getLogger(__name__)

TICKET_PARSER_PROMPT = """You are a market research data processing specialist. You parse "Reporting Tickets" — documents that specify what crosstabs and analyses should be produced from a survey dataset.

Your job: read the ticket + the variable list from the .sav file, and output a JSON spec that our tabulation engine can execute.

## What you extract

1. **banners** — The column variables for crosstabs (demographics, segments, data cuts).
   - "Total" = include a Total column (we do this automatically)
   - "PsO Patients / PsA Patients / UC-CD Patients" = custom groups based on a screener variable

2. **stubs** — The row variables (questions to tabulate). Match them to actual variable names in the dataset.
   - Ticket may say "S2" or "Q1A" — find the matching variable name in the available list.
   - Use the variable labels to match when names don't align.

3. **significance_level** — Usually "95% C.I." → 0.95. Default 0.95.

4. **nets** — T2B (Top 2 Box), B2B (Bottom 2 Box), M3B (Middle 3 Box).
   - For 5-point scales: T2B=[4,5], B2B=[1,2]
   - For 7-point scales: T2B=[6,7], B2B=[1,2], M3B=[3,4,5]
   - For 10-point scales: T2B=[9,10], B2B=[1,2]

5. **weight** — Weight variable name if mentioned.

6. **notes** — Anything you couldn't match or weren't sure about.

## Rules

- ONLY use variable names that exist in <available_variables>
- When matching, compare BOTH the variable name AND its label
- If ticket says "S2" and the .sav has a variable labeled "S2: Diagnosed Indications", that's a match
- If you can't find a match, add it to notes — don't invent variable names
- Ignore chart format instructions (Stacked Bar, Horizontal Bar, etc.) — we only generate tables
- Ignore open-end categorization instructions
- Return ONLY valid JSON, no markdown

## Output format

{
  "banners": ["variable_name_1", "variable_name_2"],
  "stubs": ["variable_name_3", "variable_name_4", ...],
  "significance_level": 0.95,
  "nets": {
    "variable_name_3": {"Top 2 Box": [4, 5], "Bottom 2 Box": [1, 2]},
    "variable_name_4": {"Top 2 Box": [4, 5], "Bottom 2 Box": [1, 2]}
  },
  "weight": null,
  "include_means": true,
  "title": "Study title from ticket",
  "notes": ["Could not match 'Q15' to any variable in the dataset"],
  "matched_variables": {
    "S2": {"matched_to": "Q_10", "label": "S2: Diagnosed Indications", "confidence": "high"},
    "Q1A": {"matched_to": "Q_20", "label": "Q1A: Treatment satisfaction", "confidence": "medium"}
  }
}"""


class TicketParser:
    """Parse Reporting Ticket .docx files using Claude Sonnet."""

    def __init__(self):
        settings = get_settings()
        if not settings.anthropic_api_key:
            raise ValueError("ANTHROPIC_API_KEY is required for ticket parsing")

        from anthropic import AsyncAnthropic
        self.client = AsyncAnthropic(api_key=settings.anthropic_api_key)

    async def parse(
        self,
        docx_bytes: bytes,
        available_variables: list[dict[str, Any]] | list[str] | None = None,
    ) -> dict[str, Any]:
        """Parse a .docx Reporting Ticket into a TabulateSpec-compatible plan.

        Args:
            docx_bytes: Raw .docx file bytes
            available_variables: List of variable names (str) or VariableInfo dicts

        Returns:
            Dict with banners, stubs, significance_level, nets, weight, notes, matched_variables
        """
        text = self._extract_text(docx_bytes)
        if not text.strip():
            return self._empty_result("Empty document — no text could be extracted")

        # Build the variable list for matching
        user_content = f"<ticket>\n{text[:10000]}\n</ticket>"

        if available_variables:
            if isinstance(available_variables[0], str):
                var_summary = "\n".join(f"- {v}" for v in available_variables[:300])
            else:
                var_summary = "\n".join(
                    f"- {v['name']}: {v.get('label', '')} (type={v.get('type', '?')}, "
                    f"labels={list(v.get('value_labels', {}).values())[:5] if v.get('value_labels') else 'none'})"
                    for v in available_variables[:300]
                )
            user_content += f"\n\n<available_variables>\n{var_summary}\n</available_variables>"

        user_content += "\n\nParse this reporting ticket. Match each question mentioned in the ticket to the closest variable in the dataset. Output ONLY the JSON spec."

        # Call Sonnet for better reasoning + variable matching
        response = await self.client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            system=TICKET_PARSER_PROMPT,
            messages=[{"role": "user", "content": user_content}],
        )

        response_text = response.content[0].text if response.content else ""
        plan = self._parse_json(response_text)

        # Ensure all required fields exist
        plan.setdefault("banners", [])
        plan.setdefault("stubs", [])
        plan.setdefault("significance_level", 0.95)
        plan.setdefault("nets", {})
        plan.setdefault("weight", None)
        plan.setdefault("include_means", True)
        plan.setdefault("notes", [])
        plan.setdefault("matched_variables", {})
        plan.setdefault("title", "")

        # Legacy compat: also provide flat fields used by auto_analyze
        if plan["banners"]:
            plan["sig_level"] = plan["significance_level"]
        if plan.get("additional_cuts"):
            plan["ticket_additional_cuts"] = plan["additional_cuts"]

        n_stubs = len(plan["stubs"])
        n_banners = len(plan["banners"])
        n_notes = len(plan["notes"])
        logger.info(
            "[TICKET] Parsed: %d banners, %d stubs, sig=%.2f, %d nets, %d notes, model=sonnet",
            n_banners, n_stubs, plan["significance_level"], len(plan["nets"]), n_notes,
        )

        return plan

    @staticmethod
    def _empty_result(note: str) -> dict:
        return {
            "banners": [], "stubs": [], "significance_level": 0.95,
            "nets": {}, "weight": None, "include_means": True,
            "notes": [note], "matched_variables": {}, "title": "",
        }

    @staticmethod
    def _parse_json(text: str) -> dict:
        """Extract JSON from response, handling markdown code blocks."""
        text = text.strip()
        # Remove markdown code block if present
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

        try:
            return json.loads(text)
        except json.JSONDecodeError:
            start = text.find("{")
            end = text.rfind("}") + 1
            if start >= 0 and end > start:
                try:
                    return json.loads(text[start:end])
                except json.JSONDecodeError:
                    pass
            return {"notes": ["Failed to parse Sonnet response as JSON"]}

    @staticmethod
    def _extract_text(docx_bytes: bytes) -> str:
        """Extract plain text from a .docx file, including tables."""
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
